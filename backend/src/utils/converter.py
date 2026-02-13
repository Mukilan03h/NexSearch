import os
import re
from typing import List, Optional
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
import markdown

from src.utils.logger import setup_logger

logger = setup_logger(__name__)

class DocumentConverter:
    """
    Converts Markdown content to PDF and Word (DOCX) formats.
    """

    @staticmethod
    def _sanitize_text(text: str) -> str:
        """
        Replace or remove characters that are not supported by standard PDF fonts (Helvetica).
        """
        # Common Unicode characters that cause issues with standard PDF fonts
        replacements = {
            "\u2014": "---", # em dash
            "\u2013": "--",  # en dash
            "\u201c": '"',   # left double quote
            "\u201d": '"',   # right double quote
            "\u2018": "'",   # left single quote
            "\u2019": "'",   # right single quote
            "\u2026": "...", # ellipsis
            "\u2212": "-",   # minus sign
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        
        # Fallback: encode to latin-1 and ignore what doesn't fit
        # This prevents the hard crash the user reported
        return text.encode("latin-1", errors="replace").decode("latin-1")

    @staticmethod
    def to_pdf(markdown_content: str, output_path: str, title: str = "Research Report") -> str:
        """
        Convert Markdown to PDF using fpdf2.
        """
        logger.info(f"Converting markdown to PDF: {output_path}")
        
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Add Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, DocumentConverter._sanitize_text(title), ln=True, align="C")
        pdf.ln(10)
        
        # Simple Markdown parsing (Headings and Paragraphs)
        lines = markdown_content.split('\n')
        
        for line in lines:
            line = DocumentConverter._sanitize_text(line)
            if not line.strip():
                pdf.ln(5)
                continue
                
            # Headings
            if line.startswith('# '):
                pdf.set_font("Helvetica", "B", 14)
                pdf.cell(0, 10, line[2:].strip(), ln=True)
            elif line.startswith('## '):
                pdf.set_font("Helvetica", "B", 13)
                pdf.cell(0, 10, line[3:].strip(), ln=True)
            elif line.startswith('### '):
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 10, line[4:].strip(), ln=True)
            # List items
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                pdf.set_font("Helvetica", "", 10)
                text = line.strip()[2:]
                pdf.multi_cell(0, 6, f"* {text}")
            # Normal text
            else:
                pdf.set_font("Helvetica", "", 10)
                # Clean basic markdown bold/italic for simplicity in PDF
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)
                pdf.multi_cell(0, 6, clean_text)
        
        pdf.output(output_path)
        return output_path

    @staticmethod
    def to_docx(markdown_content: str, output_path: str, title: str = "Research Report") -> str:
        """
        Convert Markdown to Word using python-docx.
        """
        logger.info(f"Converting markdown to DOCX: {output_path}")
        
        doc = Document()
        doc.add_heading(title, 0)
        
        lines = markdown_content.split('\n')
        for line in lines:
            if not line.strip():
                continue
                
            if line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:].strip(), level=3)
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                doc.add_paragraph(line.strip()[2:], style='List Bullet')
            else:
                # Add paragraph with basic bold/italic support
                p = doc.add_paragraph()
                
                # Simple regex for bold text
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
        
        doc.save(output_path)
        return output_path
